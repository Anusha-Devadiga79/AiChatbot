from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime
import json

def generate_pdf_export(chats: list, user_info: dict) -> bytes:
    """Generate a PDF export of chat history with analysis details."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                  fontSize=18, spaceAfter=12, textColor=colors.HexColor('#2563eb'))
    story.append(Paragraph("AI Public Health Chatbot - Chat History & Analysis", title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # User info
    story.append(Paragraph(f"User: {user_info.get('username', 'N/A')}", styles['Normal']))
    story.append(Paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 0.3 * inch))
    
    # Chat entries
    for idx, chat in enumerate(chats):
        # User message
        user_style = ParagraphStyle('UserMsg', parent=styles['Normal'],
                                     backColor=colors.HexColor('#dbeafe'),
                                     borderPadding=8, fontSize=10)
        story.append(Paragraph(f"<b>You:</b> {chat.get('message', '')}", user_style))
        story.append(Spacer(1, 0.1 * inch))
        
        # Analysis details if available
        analysis = chat.get('analysis', {})
        if isinstance(analysis, str):
            try:
                analysis = json.loads(analysis)
            except:
                analysis = {}
        
        if analysis and analysis.get('matched'):
            analysis_style = ParagraphStyle('Analysis', parent=styles['Normal'],
                                           backColor=colors.HexColor('#f0fdf4'),
                                           borderPadding=8, fontSize=9)
            
            analysis_text = f"""
            <b>Disease:</b> {analysis.get('disease', 'N/A')}<br/>
            <b>Severity:</b> {analysis.get('severity', 'N/A')}<br/>
            <b>Confidence:</b> {int(analysis.get('confidence', 0) * 100)}%<br/>
            <b>Description:</b> {analysis.get('description', 'N/A')}<br/>
            <b>Prevention:</b> {analysis.get('prevention', 'N/A')}<br/>
            <b>When to See Doctor:</b> {analysis.get('when_to_see_doctor', 'N/A')}
            """
            story.append(Paragraph(analysis_text, analysis_style))
            story.append(Spacer(1, 0.1 * inch))
        
        # Bot response
        bot_style = ParagraphStyle('BotMsg', parent=styles['Normal'],
                                    backColor=colors.HexColor('#fef3c7'),
                                    borderPadding=8, fontSize=10)
        response_text = chat.get('response', '').replace('\n', '<br/>')
        story.append(Paragraph(f"<b>Assistant:</b> {response_text}", bot_style))
        
        # Timestamp
        ts = chat.get('timestamp', '')
        if ts:
            story.append(Paragraph(f"<i>{ts}</i>", styles['Normal']))
        
        # Uploaded files info
        uploaded_files = chat.get('uploaded_files')
        if uploaded_files:
            files_list = uploaded_files if isinstance(uploaded_files, list) else uploaded_files.split(',')
            files_text = ", ".join([os.path.basename(f) for f in files_list])
            story.append(Paragraph(f"<i>📎 Uploaded: {files_text}</i>", styles['Normal']))
        
        story.append(Spacer(1, 0.3 * inch))
        
        # Page break every 5 chats
        if (idx + 1) % 5 == 0:
            story.append(PageBreak())
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_doc_export(chats: list, user_info: dict) -> bytes:
    """Generate a DOCX export of chat history with analysis details."""
    doc = Document()
    
    # Title
    title = doc.add_heading("AI Public Health Chatbot - Chat History & Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"User: {user_info.get('username', 'N/A')}")
    doc.add_paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    
    for chat in chats:
        # User message
        p = doc.add_paragraph()
        p.add_run("You: ").bold = True
        p.add_run(chat.get('message', ''))
        
        # Analysis details if available
        analysis = chat.get('analysis', {})
        if isinstance(analysis, str):
            try:
                analysis = json.loads(analysis)
            except:
                analysis = {}
        
        if analysis and analysis.get('matched'):
            doc.add_heading("Analysis Results", level=3)
            
            analysis_table = doc.add_table(rows=7, cols=2)
            analysis_table.style = 'Light Grid Accent 1'
            
            rows = analysis_table.rows
            rows[0].cells[0].text = "Disease"
            rows[0].cells[1].text = analysis.get('disease', 'N/A')
            rows[1].cells[0].text = "Severity"
            rows[1].cells[1].text = analysis.get('severity', 'N/A')
            rows[2].cells[0].text = "Confidence"
            rows[2].cells[1].text = f"{int(analysis.get('confidence', 0) * 100)}%"
            rows[3].cells[0].text = "Description"
            rows[3].cells[1].text = analysis.get('description', 'N/A')
            rows[4].cells[0].text = "Prevention"
            rows[4].cells[1].text = analysis.get('prevention', 'N/A')
            rows[5].cells[0].text = "When to See Doctor"
            rows[5].cells[1].text = analysis.get('when_to_see_doctor', 'N/A')
            rows[6].cells[0].text = "Other Conditions"
            rows[6].cells[1].text = ", ".join(analysis.get('other_possible_conditions', []))
        
        # Bot response
        p2 = doc.add_paragraph()
        p2.add_run("Assistant: ").bold = True
        p2.add_run(chat.get('response', ''))
        
        # Timestamp
        ts = chat.get('timestamp', '')
        if ts:
            p_ts = doc.add_paragraph(ts)
            p_ts.italic = True
        
        # Uploaded files info
        uploaded_files = chat.get('uploaded_files')
        if uploaded_files:
            files_list = uploaded_files if isinstance(uploaded_files, list) else uploaded_files.split(',')
            files_text = ", ".join([os.path.basename(f) for f in files_list])
            doc.add_paragraph(f"📎 Uploaded: {files_text}")
        
        doc.add_paragraph("─" * 80)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_json_export(chats: list, user_info: dict) -> bytes:
    """Generate a JSON export of chat history with full analysis data."""
    export_data = {
        "user": user_info,
        "exported_at": datetime.now().isoformat(),
        "total_chats": len(chats),
        "chats": []
    }
    
    for chat in chats:
        analysis = chat.get('analysis', {})
        if isinstance(analysis, str):
            try:
                analysis = json.loads(analysis)
            except:
                analysis = {}
        
        uploaded_files = chat.get('uploaded_files')
        files_list = []
        if uploaded_files:
            files_list = uploaded_files if isinstance(uploaded_files, list) else uploaded_files.split(',')
            files_list = [os.path.basename(f) for f in files_list]
        
        chat_data = {
            "chat_id": chat.get("chat_id"),
            "timestamp": chat.get("timestamp"),
            "message": chat.get("message"),
            "response": chat.get("response"),
            "uploaded_files": files_list,
            "analysis": analysis
        }
        export_data["chats"].append(chat_data)
    
    json_str = json.dumps(export_data, indent=2, ensure_ascii=False)
    return json_str.encode('utf-8')

def generate_csv_export(chats: list, user_info: dict) -> bytes:
    """Generate a CSV export of chat history with key analysis fields."""
    import csv
    from io import StringIO
    
    output = StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow([
        "Chat ID", "Timestamp", "User Message", "Disease", "Severity", 
        "Confidence", "Symptoms", "Prevention", "When to See Doctor", 
        "Uploaded Files", "Response"
    ])
    
    # Data rows
    for chat in chats:
        analysis = chat.get('analysis', {})
        if isinstance(analysis, str):
            try:
                analysis = json.loads(analysis)
            except:
                analysis = {}
        
        uploaded_files = chat.get('uploaded_files')
        files_str = ""
        if uploaded_files:
            files_list = uploaded_files if isinstance(uploaded_files, list) else uploaded_files.split(',')
            files_str = "; ".join([os.path.basename(f) for f in files_list])
        
        symptoms_str = "; ".join(analysis.get('symptoms', [])) if analysis.get('symptoms') else ""
        
        writer.writerow([
            chat.get("chat_id", ""),
            chat.get("timestamp", ""),
            chat.get("message", ""),
            analysis.get("disease", "") if analysis.get("matched") else "",
            analysis.get("severity", "") if analysis.get("matched") else "",
            f"{int(analysis.get('confidence', 0) * 100)}%" if analysis.get("confidence") else "",
            symptoms_str,
            analysis.get("prevention", "") if analysis.get("matched") else "",
            analysis.get("when_to_see_doctor", "") if analysis.get("matched") else "",
            files_str,
            chat.get("response", "")[:200] + "..." if len(chat.get("response", "")) > 200 else chat.get("response", "")
        ])
    
    csv_bytes = output.getvalue().encode('utf-8')
    output.close()
    return csv_bytes

def _strip_html(text: str) -> str:
    """Remove basic HTML tags for plain-text rendering in PDF/DOCX."""
    import re
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
    return text.strip()


def generate_report_pdf(summary: str = "", previous: str = "", current: str = "",
                        language: str = "English", prev_name: str = "Previous Report",
                        curr_name: str = "Current Report") -> bytes:
    """Generate a PDF containing only the AI Health Summary."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=60, leftMargin=60,
                            topMargin=60, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title_style = ParagraphStyle('ReportTitle', parent=styles['Heading1'],
                                  fontSize=20, spaceAfter=4,
                                  textColor=colors.HexColor('#1d4ed8'))
    story.append(Paragraph("AI Health Summary", title_style))

    meta_style = ParagraphStyle('Meta', parent=styles['Normal'],
                                 fontSize=9, textColor=colors.HexColor('#6b7280'), spaceAfter=16)
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}  |  Language: {language}",
        meta_style
    ))

    # Summary content
    summary_style = ParagraphStyle('Summary', parent=styles['Normal'],
                                    fontSize=11, leading=17,
                                    backColor=colors.HexColor('#eff6ff'),
                                    borderPadding=12,
                                    spaceAfter=6)
    for line in _strip_html(summary).split('\n'):
        line = line.strip()
        if line:
            story.append(Paragraph(line, summary_style))

    story.append(Spacer(1, 0.25 * inch))

    # Disclaimer
    disc_style = ParagraphStyle('Disc', parent=styles['Normal'],
                                 fontSize=8, textColor=colors.HexColor('#dc2626'),
                                 backColor=colors.HexColor('#fef2f2'),
                                 borderPadding=6)
    story.append(Paragraph(
        "⚠️ This summary is for awareness purposes only. "
        "Always consult a qualified healthcare professional for proper diagnosis and treatment.",
        disc_style
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_report_doc(summary: str = "", previous: str = "", current: str = "",
                        language: str = "English", prev_name: str = "Previous Report",
                        curr_name: str = "Current Report") -> bytes:
    """Generate a DOCX containing only the AI Health Summary."""
    doc = Document()

    # Title
    title = doc.add_heading("AI Health Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}   |   Language: {language}"
    )
    if meta.runs:
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    doc.add_paragraph("")

    # Summary content
    h = doc.add_heading("📈 Health Summary", level=2)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(30, 64, 175)

    for line in _strip_html(summary).split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph("")

    # Disclaimer
    disc = doc.add_paragraph(
        "⚠️ This summary is for awareness purposes only. "
        "Always consult a qualified healthcare professional for proper diagnosis and treatment."
    )
    if disc.runs:
        disc.runs[0].font.color.rgb = RGBColor(220, 38, 38)
        disc.runs[0].font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
